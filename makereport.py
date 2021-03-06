#!/usr/bin/python

import os
import sys
import time
import subprocess
import elementtree
from docx import Document
from docx.shared import Inches
import parsenessus
import parsenmap
from host import Host
from service import Service
from vulnerability import Vulnerability
import glob

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH

#init variables
document = Document()
hosts = []
path = os.environ['HOME']+'/Desktop/audit-scans/'


def main():
	os.system('clear')
	print("Welcome to the Audit Report Builder")
	print("Run this after completing the nmapscan and nessusscan programs")

	date = time.strftime("%d/%m/%Y")
	school_name = raw_input("School name: ")

	#read all files from the folder that start with the date(user input)
	#show suggestions
	allfiles = glob.glob(path+"*.xml")
	datelist = []
	for file in allfiles:
		datefile = file.replace(path,'')
		groups = datefile.split('-')
		date = '-'.join(groups[:3])
		datelist.append(date)

	#remove duplicated
	unique_datelist = []
	for date in datelist:
		if not (date in unique_datelist):
			unique_datelist.append(date)

	print ("Available dates are: ")
	for date in unique_datelist:
		print ("\t" + str(date))


	#check if good date
	valid_date = False
	while not (valid_date):
		audit_date = raw_input('Audit date (in format: yyyy-mm-dd): ')
		if not (audit_date in unique_datelist):
			print 'not a valid date, see available dates list'
			valid_date = False
		else:
			valid_date = True

	#add nmap hosts
	nmaphosts = add_hosts('nmap',audit_date)

	if( len(nmaphosts) == 0):
		print('no nmaphosts found, exiting')
		exit()

	#add nessus hosts
	nessushosts = add_hosts('nessus',audit_date)

	#fuse hosts info together
	for nmaphost in nmaphosts:
		current = find_host(nessushosts,nmaphost.ip)
		if current is not None:
			if (nmaphost.mac is None): nmaphost.mac = current.mac
			if (nmaphost.os is None): nmaphost.os = current.os
			if (nmaphost.hostname is None): nmaphost.hostname = current.hostname
			nmaphost.vulnerabilities = current.vulnerabilities
			hosts.append(nmaphost)
		else:
			hosts.append(nmaphost)


	#select hosts
	print('All hosts in database:')
	time.sleep(1)
	for host in hosts:
		print ('\t' + str(host.ip))


	selected_hosts_ok = False
	while not (selected_hosts_ok):
		#check for criteria end display new list of hosts
		nrOfServices = int(raw_input("Minimum number of detected services? (nmap): "))
		severity = float(raw_input("Vulnerability severity should be at least? (0.0 - 10.0): "))

		determined_hosts = determine_hosts(hosts, nrOfServices,severity)

		print(str(len(determined_hosts)) + ' hosts match those criteria:')
		for host in determined_hosts:
			print(str(host.ip))

		select_ok = raw_input("Is this selection ok? (y/n): ")
		if (select_ok == "y"): selected_hosts_ok = True
		else: selected_hosts_ok = False

	print ("Enter the hosts you want to display in the report, 'all' to select all hosts, or 'stop' to stop")
	selected_hosts = []
	maxLength = len(determined_hosts)
	while len(selected_hosts) < maxLength:
		item = raw_input("Host: ")
		if (item == 'stop'):
			break
		if(item == 'all'):
			selected_hosts = determined_hosts
			break
		else:
			selected_hosts.append(find_host(hosts,item))

	print (str(len(selected_hosts)) + " hosts selected")


	#build front page
	howest_img_p = document.add_paragraph()
	howest_img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	img_r = howest_img_p.add_run()
	img_r.add_picture('howest.jpg', width=Inches(4.9))

	front_page1 = document.add_paragraph("Computer & CyberCrime Professional\n\n")
	front_page1.alignment = WD_ALIGN_PARAGRAPH.CENTER
	front_page1.style = document.styles['Heading 5']

	front_page = document.add_paragraph("Audit report for:")
	front_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
	front_page.style = document.styles['Heading 2']

	front_page0 = document.add_paragraph(str(school_name))
	front_page0.alignment = WD_ALIGN_PARAGRAPH.CENTER
	front_page0.style = document.styles['Heading 1']


	front_page2 = document.add_paragraph("Conducted on: " +str(audit_date))
	front_page2.alignment = WD_ALIGN_PARAGRAPH.CENTER
	front_page2.style = document.styles['Heading 3']

	document.add_paragraph("\n\n\n\n\n\n\n\n\n\n\n\n\n")

	front_page3 = document.add_paragraph("Report generated with: \n http://github.com/ThimoDeSouter/security-audit-scripts")
	front_page3.alignment = WD_ALIGN_PARAGRAPH.CENTER
	front_page3.style = document.styles['Heading 6']

	document.add_page_break()
	#end build front page

	build_table(selected_hosts,severity)

	document.save(os.environ['HOME'] + '/Desktop/audit-report-' + school_name + '.docx')
	print 'audit-report-'+school_name+ '.docx saved in' + os.environ['HOME']+'/Desktop'
	return

def displayNiktoProcess(exitlist):
	os.system('clear')
	print ('Running Nikto scans...\n')
	print ('\033[1m' +'HOST \t\t ' + '\033[1m'+ 'STATUS')
	for ip,code in exitlist.items():
		status = "x"
		if(code is None): status = "Working"
		elif(code is 0): status = "Done"
		print '\033[0m' + ip + " \t " + status

	count = 20
	while(count >=0):
    		sys.stdout.write("\rRefresh in:{0}>>".format(count))
    		sys.stdout.flush()
		count -=1
    		time.sleep(1)



def find_host(hosts,ip):
	foundHost = None
	for host in hosts:
		if(host.ip == ip):
			foundHost = host
	return foundHost


def add_hosts(type,audit_date):
	filetype = None
	if(type) == "nmap":
		filetype = ".xml"
	if(type) == "nessus":
		filetype = ".nessus"

	if not os.path.exists(path):
		print 'Please run nmapscan or nessusscan first'
		exit()

	scans = []
	for file in os.listdir(path):
		if os.path.isfile(os.path.join(path,file)) and file.startswith(audit_date) and file.endswith(filetype):
			scans.append(file)

	print 'found ' + str(len(scans)) + ' ' + str(type) + ' scans'

	print ('processing...')

	hosts = []
	for scan in scans:
		file_path = str(path) +str(scan)

		if(type == "nmap"):
			hostslocal = parsenmap.main(file_path)
		if(type =="nessus"):
			hostslocal = parsenessus.main(file_path)

		for hostlocal in hostslocal:
			duplicate = None
			duplicatehost = None
			for host in hosts:
				if host.ip == hostlocal.ip:
					duplicate = True
					duplicatehost = host
			if(duplicate):
				if( len(hostlocal.services) > len(duplicatehost.services) ):
					hosts.remove(duplicatehost)
					hosts.append(hostlocal)
			else:
				hosts.append(hostlocal)

	return hosts

def determine_hosts(hosts, nrOfServices, severity):
	myhosts = []
	for host in hosts:
		addhost = False
		if( (host is not None) and (len(host.services) >= nrOfServices)):
			for vuln in host.vulnerabilities:
				if float(vuln.base_score) >= float(severity):
					addhost = True

		if(addhost):
			myhosts.append(host)
	return myhosts

def build_table(selected_hosts,severity):

	#select hosts
	myhosts = selected_hosts


	#nikto
	web_hosts = []
	print("Searching for hosts that run a webserver...")
	#loop through all hosts and check if port 80 or 443 is open
	addhost = False
	for host in myhosts:
		for service in host.services:
			if (service.port == '80' or service.port == '443' ):
				addhost = True

		if(addhost):
			web_hosts.append(host)

	print('found ' + str(len(web_hosts))+' webhosts:')
	for host in web_hosts:
		print ('\t' + str(host.ip))

	print ("You can gather information about the webserver(s) by running nikto.\nNote that you have to be connected to the local network for this to work.")
	run_nikto = raw_input("Run nikto on these hosts? (y/n): ")
	if(run_nikto =="y"):

		hostlist = []
		exits = []

		processes = []
		for host in web_hosts:
			tmpfile = os.tmpfile()
			command="nikto -h " + host.ip
			proc = subprocess.Popen([command],stdout=tmpfile, shell=True)
			processes.append((proc,tmpfile,host.ip))
			hostlist.append(host.ip)
			exits.append(proc.poll())

		exitlist = dict(zip(hostlist, exits))

		while None in exitlist.values():
			for proc,file,ip in processes:
				exitlist[ip] = proc.poll()
			displayNiktoProcess(exitlist)


		for proc,file, ip in processes:
				file.seek(0)
				nikto_out = file.read()
				file.close()
				nikhost = find_host(selected_hosts,ip)
				nikhost.nikto = nikto_out


	#end nikto

	print ('\nGenerating report...')
	hostcount_total = len(myhosts)
	count = 0
	for host in myhosts:
		#progress
		count +=1
		sys.stdout.write("\rProcessed host({0} of ".format(count))
		sys.stdout.flush()
		print(str(hostcount_total)+")")

		services = host.services
		vulnerabilities = host.vulnerabilities
		rowcount = len(services)

		#main table
		main_table = document.add_table(rows=0, cols=1)
		main_table.autofit = False
		main_table.style = 'Table Grid'

		#info section
		info_section = main_table.add_row()

		p_info = info_section.cells[0].paragraphs[0]
		p_info.add_run('IP: ').bold = True
		p_info.add_run(str(host.ip)+'\n')

		p_info.add_run('Hostname: ').bold = True
		p_info.add_run(str(host.hostname)+'\n')

		p_info.add_run('Operating System: ').bold = True
		p_info.add_run(str(host.os)+'\n')

		p_info.add_run('MAC Address: ').bold = True
		p_info.add_run(str(host.mac))


		#open services text section
		open_services_section = main_table.add_row()
		open_services_p = open_services_section.cells[0].paragraphs[0]
		open_services_p.add_run('Detected Open Services:').bold = True

		#services section
		services_section = main_table.add_row()
		remove_p = services_section.cells[0].paragraphs[0]
		p = remove_p._element
		p.getparent().remove(p)
		p._p = p._element = None
		services_table = services_section.cells[0].add_table(rows=0, cols=4)
		services_table.style = 'Table Grid'
		services_table.columns[0].width = Inches(1.5)
		services_table.columns[1].width = Inches(1)
		services_table.columns[2].width = Inches(1)
		services_table.columns[3].width = Inches(2)



		#header row services
		header_row = services_table.add_row()

		name_header = header_row.cells[0]
		name_header.text = "Name"
		name_run = name_header.paragraphs[0].runs[0]
		name_run.font.bold = True

		port_header = header_row.cells[1]
		port_header.text = "Port"
		port_run = port_header.paragraphs[0].runs[0]
		port_run.font.bold = True

		protocol_header = header_row.cells[2]
		protocol_header.text = "Protocol"
		protocol_run = protocol_header.paragraphs[0].runs[0]
		protocol_run.font.bold = True

		info_header = header_row.cells[3]
		info_header.text = "Info"
		info_run = info_header.paragraphs[0].runs[0]
		info_run.font.bold = True

		for service in services:
			row = services_table.add_row()
			row.cells[0].text = service.name
			row.cells[1].text = service.port
			row.cells[2].text = service.protocol
			row.cells[3].text = str(service.product) + ' ' + str(service.extrainfo) + ' ' + str(service.version)
			row.cells[3].text = row.cells[3].text.replace('None',' ')


		if ( (len(vulnerabilities) >= 1)):
			vulnprint = False
			for vuln in vulnerabilities:
				if( float(vuln.base_score) >= float(severity)):
					vulnprint = True

			if(vulnprint):
				#vulnerabilities text section
				vulns_section = main_table.add_row()
				vulns_p = vulns_section.cells[0].paragraphs[0]
				vulns_p.add_run('Detected Vulnerabilities:').bold = True


				#vulnerabilities section
				vulnerabilities_section = main_table.add_row()
				vuln_p_rem = vulnerabilities_section.cells[0].paragraphs[0]
				p2 = vuln_p_rem._element
				p2.getparent().remove(p2)
				p2._p2 = p2._element = None

				vulnerabilities_table = vulnerabilities_section.cells[0].add_table(rows=0,cols=2)
				vulnerabilities_table.autofit = False
				vulnerabilities_table.columns[0].width = Inches(1)
				vulnerabilities_table.columns[1].width = Inches(4.5)
				vulnerabilities_table.style = 'Table Grid'

				#header row vulns
				header_row = vulnerabilities_table.add_row()

				severity_header = header_row.cells[0]
				severity_header.text = "Severity"
				severity_run = severity_header.paragraphs[0].runs[0]
				severity_run.font.bold = True

				description_header = header_row.cells[1]
				description_header.text = "Description"
				description_run = description_header.paragraphs[0].runs[0]
				description_run.font.bold = True


				#actual vuls
				vulnerabilities.sort(key=lambda x: float(x.base_score), reverse=True)
				for vulnerability in vulnerabilities:
					if ( float(vulnerability.base_score) >= severity):
						row = vulnerabilities_table.add_row()
						row.cells[0].text = str(vulnerability.base_score)

						sev = float(vulnerability.base_score)

						#set vuln color
						if( sev >= 9.5): #critical
							shading_elm = parse_xml(r'<w:shd {} w:fill="ff0000"/>'.format(nsdecls('w')))
							row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)

						elif( sev <= 9.4 and sev >=7.6): #high
							shading_elm = parse_xml(r'<w:shd {} w:fill="ee7600"/>'.format(nsdecls('w')))
							row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)

						elif( sev <= 7.5 and sev >=4.0): #medium
							shading_elm = parse_xml(r'<w:shd {} w:fill="ffa500"/>'.format(nsdecls('w')))
							row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)

						elif( sev <=3.9 and sev >=1.1): #low
							shading_elm = parse_xml(r'<w:shd {} w:fill="00ff00"/>'.format(nsdecls('w')))
							row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)

						elif( sev <=1.0): #info
							shading_elm = parse_xml(r'<w:shd {} w:fill="0000ff"/>'.format(nsdecls('w')))
							row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)



						#p = row.cells[1].add_paragraph()
						p = row.cells[1].paragraphs[0]
						p.add_run('Name: ').bold = True
						p.add_run(str(vulnerability.name)+'\n')

						p.add_run('Synopsis: ').bold = True
						p.add_run(str(vulnerability.synopsis)+'\n')

						p.add_run('CVE: ').bold = True
						p.add_run(str(vulnerability.cve)+'\n')

						p.add_run('CVE Date: ').bold = True
						p.add_run(str(vulnerability.pub_date)+'\n')

						p.add_run('Score: ').bold = True
						p.add_run(str(vulnerability.severity)+'\n')

						p.add_run('Solution: ').bold = True
						p.add_run(str(vulnerability.solution))


		#nikto section
		if host.nikto is not None:
			nikto_section = main_table.add_row()
			p_nikto = nikto_section.cells[0].paragraphs[0]
			p_nikto.add_run(str(host.nikto))
		

		document.add_paragraph("")
	return

if __name__ == "__main__":
	main()
